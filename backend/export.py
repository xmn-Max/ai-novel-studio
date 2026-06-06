from __future__ import annotations

import io
from typing import Optional

import yaml
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ExportError(Exception):
    pass


def _parse_yaml(yaml_str: str) -> dict:
    try:
        data = yaml.safe_load(yaml_str)
        if not isinstance(data, dict):
            raise ExportError("无效的 YAML 格式")
        return data
    except yaml.YAMLError as e:
        raise ExportError(f"YAML 解析失败: {e}")


def export_pdf(yaml_str: str) -> bytes:
    data = _parse_yaml(yaml_str)
    meta = data.get("meta", {})
    script = data.get("script", [])

    pdf = FPDF()
    pdf.add_page()

    font_paths = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    font_loaded = False
    for fp in font_paths:
        try:
            pdf.add_font("CN", "", fp, uni=True)
            pdf.add_font("CN", "B", fp, uni=True)
            font_loaded = True
            break
        except Exception:
            continue

    if not font_loaded:
        try:
            pdf.add_font("CN", "", "helvetica")
            pdf.add_font("CN", "B", "helvetica")
        except Exception:
            pass

    def safe_text(text: str) -> str:
        if not text:
            return ""
        try:
            text.encode("latin-1")
            return text
        except UnicodeEncodeError:
            return "".join(c if ord(c) < 256 else "?" for c in text)

    title = meta.get("title", "未命名")
    pdf.set_font("CN", "B", 16)
    pdf.cell(0, 10, safe_text(title), ln=True, align="C")
    pdf.ln(5)

    pdf.set_font("CN", "", 10)
    genre = meta.get("genre", "")
    chapters = meta.get("source_chapters", 0)
    scenes = meta.get("total_scenes", 0)
    pdf.cell(0, 6, safe_text(f"类型: {genre}  章节: {chapters}  场景: {scenes}"), ln=True, align="C")
    pdf.ln(8)

    characters = meta.get("characters", [])
    if characters:
        pdf.set_font("CN", "B", 11)
        pdf.cell(0, 7, safe_text("角色列表"), ln=True)
        pdf.set_font("CN", "", 10)
        pdf.cell(0, 6, safe_text("、".join(characters)), ln=True)
        pdf.ln(5)

    if not isinstance(script, list):
        script = []

    for i, scene in enumerate(script):
        heading = scene.get("scene_heading", f"场景 {i + 1}")
        location = scene.get("location", "")
        time_of_day = scene.get("time_of_day", "")

        pdf.set_font("CN", "B", 11)
        pdf.cell(0, 7, safe_text(heading), ln=True)

        if location or time_of_day:
            pdf.set_font("CN", "", 9)
            info_parts = []
            if location:
                info_parts.append(f"地点: {location}")
            if time_of_day:
                info_parts.append(f"时间: {time_of_day}")
            pdf.cell(0, 5, safe_text("  ".join(info_parts)), ln=True)
        pdf.ln(2)

        actions = scene.get("action", [])
        if actions:
            pdf.set_font("CN", "", 10)
            for action in actions:
                pdf.multi_cell(0, 5, safe_text(action))
            pdf.ln(2)

        dialogues = scene.get("dialogues", [])
        if dialogues:
            for d in dialogues:
                char = d.get("character", "")
                line = d.get("line", "")
                parenthetical = d.get("parenthetical", "")
                pdf.set_font("CN", "B", 10)
                pdf.cell(0, 5, safe_text(char), ln=True)
                if parenthetical:
                    pdf.set_font("CN", "", 9)
                    pdf.cell(0, 5, safe_text(f"({parenthetical})"), ln=True)
                pdf.set_font("CN", "", 10)
                pdf.multi_cell(0, 5, safe_text(f'"{line}"'))
            pdf.ln(2)

        transition = scene.get("transition", "")
        if transition:
            pdf.set_font("CN", "B", 9)
            pdf.cell(0, 5, safe_text(f"[{transition}]"), ln=True, align="R")
        pdf.ln(3)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def export_docx(yaml_str: str) -> bytes:
    data = _parse_yaml(yaml_str)
    meta = data.get("meta", {})
    script = data.get("script", [])

    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Microsoft YaHei"

    title = meta.get("title", "未命名")
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    genre = meta.get("genre", "")
    chapters = meta.get("source_chapters", 0)
    scenes = meta.get("total_scenes", 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"类型: {genre}  章节: {chapters}  场景: {scenes}").font.size = Pt(10)

    characters = meta.get("characters", [])
    if characters:
        doc.add_heading("角色列表", level=2)
        doc.add_paragraph("、".join(characters))

    if not isinstance(script, list):
        script = []

    for i, scene in enumerate(script):
        heading = scene.get("scene_heading", f"场景 {i + 1}")
        doc.add_heading(heading, level=2)

        location = scene.get("location", "")
        time_of_day = scene.get("time_of_day", "")
        if location or time_of_day:
            parts = []
            if location:
                parts.append(f"地点: {location}")
            if time_of_day:
                parts.append(f"时间: {time_of_day}")
            doc.add_paragraph("  ".join(parts)).runs[0].font.size = Pt(9)

        actions = scene.get("action", [])
        for action in actions:
            doc.add_paragraph(action, style="Body Text")

        dialogues = scene.get("dialogues", [])
        for d in dialogues:
            char = d.get("character", "")
            line = d.get("line", "")
            parenthetical = d.get("parenthetical", "")

            p = doc.add_paragraph()
            run = p.add_run(char)
            run.bold = True
            run.font.size = Pt(11)

            if parenthetical:
                p = doc.add_paragraph()
                run = p.add_run(f"({parenthetical})")
                run.italic = True
                run.font.size = Pt(9)

            p = doc.add_paragraph()
            run = p.add_run(f'"{line}"')
            run.font.size = Pt(11)

        transition = scene.get("transition", "")
        if transition:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"[{transition}]")
            run.bold = True
            run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
