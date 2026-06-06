from __future__ import annotations

import asyncio
import json
import os
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

from fastapi import FastAPI, HTTPException, Depends, Header, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

from models import ConvertRequest, ConvertResponse
from pipeline import Pipeline
from auth import register_user, login_user, verify_token, change_password
from auth import get_user_genres, add_user_genre, update_user_genre, delete_user_genre
from database import (
    init_db,
    db_get_system_genres,
    db_get_user_genres as db_user_genres,
    db_create_conversion,
    db_update_conversion_progress,
    db_update_conversion_status,
    db_save_conversion_result,
    db_get_conversion,
    db_list_conversions,
    db_delete_conversion,
    db_update_conversion_yaml,
    db_mark_stale_conversions,
)
from middleware import register_exception_handlers
from export import export_pdf, export_docx, ExportError

init_db()

limiter = Limiter(key_func=get_remote_address, default_limits=["60/minute"])

app = FastAPI(title="AI Novel Studio API", version="0.4.0")
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

register_exception_handlers(app)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")


def _load_genres() -> list[dict[str, Any]]:
    genres = db_get_system_genres()
    if not genres:
        return [{"name": "叙事", "guidance": "这是一部叙事小说。请重点关注故事结构、人物心理。", "keywords": []}]
    return [{"name": g["name"], "guidance": g.get("guidance", ""), "keywords": g.get("keywords", [])} for g in genres]


def _merged_genres(username: str) -> list[dict[str, Any]]:
    system = _load_genres()
    user = db_user_genres(username)
    result: list[dict[str, Any]] = []
    for g in system:
        result.append({**g, "readonly": True})
    for g in user:
        result.append({**g, "readonly": False})
    return result


def _get_pipeline(genre: str = "叙事", title: str = "") -> Pipeline:
    if not API_KEY:
        raise RuntimeError("DEEPSEEK_API_KEY 环境变量未设置")
    return Pipeline(api_key=API_KEY, genre=genre, title=title)


async def _run_pipeline(task_id: str, text: str, genre: str, title: str, username: str) -> None:
    print(f"[DEBUG] _run_pipeline START task_id={task_id[:8]}... text_len={len(text)} genre={genre}", flush=True)
    with open(os.path.join(os.path.dirname(__file__), "debug.log"), "a", encoding="utf-8") as f:
        f.write(f"START task_id={task_id[:8]}... len={len(text)}\n")

    db_update_conversion_status(task_id, "processing")

    async def progress_callback(step: int, total: int, step_name: str, message: str) -> None:
        print(f"[DEBUG] progress_callback step={step}/{total} {step_name}: {message}", flush=True)
        db_update_conversion_progress(task_id, {
            "step": step,
            "total": total,
            "step_name": step_name,
            "message": message,
        })

    try:
        pipeline = _get_pipeline(genre, title)
        result = await pipeline.run(text, progress_callback)

        intermediate = result.pop("_intermediate", None)
        intermediate["original_text"] = text
        intermediate["genre"] = genre
        intermediate["title"] = title

        yaml_output = result.get("yaml", "")
        meta_json = json.dumps(result.get("meta", {}), ensure_ascii=False)
        intermediate_json = json.dumps(intermediate, ensure_ascii=False)

        db_save_conversion_result(
            task_id,
            yaml_output=yaml_output,
            meta_json=meta_json,
            intermediate_json=intermediate_json,
            cleaned_text=intermediate.get("cleaned_text", ""),
        )
    except Exception as e:
        print(f"[DEBUG] _run_pipeline FAILED: {type(e).__name__}: {e}")
        db_update_conversion_status(task_id, "failed", str(e))


async def _cleanup_stale_tasks(interval_seconds: int = 600) -> None:
    while True:
        await asyncio.sleep(interval_seconds)
        try:
            db_mark_stale_conversions(minutes=30)
        except Exception:
            pass


def _require_auth(authorization: str = Header(default="")) -> dict:
    token = ""
    if authorization.startswith("Bearer "):
        token = authorization[7:]
    user = verify_token(token)
    if not user:
        raise HTTPException(status_code=401, detail="未登录或登录已过期")
    return user


def _check_auth(token: str = "") -> dict:
    user = verify_token(token)
    if not user:
        raise HTTPException(status_code=401, detail="token 无效")
    return user


@app.on_event("startup")
async def startup():
    asyncio.create_task(_cleanup_stale_tasks())


# --- Auth ---

class AuthRequest(BaseModel):
    username: str
    password: str


class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str


@app.post("/api/auth/register")
@limiter.limit("10/minute")
async def api_register(request: Request, req: AuthRequest) -> dict[str, str]:
    if not req.username or not req.password:
        raise HTTPException(status_code=400, detail="用户名和密码不能为空")
    if len(req.username) < 2 or len(req.username) > 20:
        raise HTTPException(status_code=400, detail="用户名长度需在 2-20 之间")
    if len(req.password) < 4:
        raise HTTPException(status_code=400, detail="密码至少需要 4 个字符")
    try:
        return register_user(req.username.strip(), req.password)
    except ValueError as e:
        raise HTTPException(status_code=409, detail=str(e))


@app.post("/api/auth/login")
@limiter.limit("20/minute")
async def api_login(request: Request, req: AuthRequest) -> dict[str, str]:
    if not req.username or not req.password:
        raise HTTPException(status_code=400, detail="用户名和密码不能为空")
    result = login_user(req.username.strip(), req.password)
    if not result:
        raise HTTPException(status_code=401, detail="用户名或密码错误")
    return result


@app.get("/api/auth/me")
async def api_me(user: dict = Depends(_require_auth)) -> dict:
    return {"username": user["username"], "logged_in": True}


@app.post("/api/auth/change-password")
async def api_change_password(req: ChangePasswordRequest, user: dict = Depends(_require_auth)) -> dict[str, str]:
    if not req.old_password or not req.new_password:
        raise HTTPException(status_code=400, detail="原密码和新密码不能为空")
    try:
        change_password(user["username"], req.old_password, req.new_password)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return {"detail": "密码修改成功"}


# --- Genre CRUD (per-user) ---

class GenreItem(BaseModel):
    name: str
    guidance: str = ""
    keywords: list[str] = []


@app.get("/api/genres")
async def list_genres(user: dict = Depends(_require_auth)) -> list[dict[str, Any]]:
    return _merged_genres(user["username"])


@app.post("/api/genres")
async def add_genre(item: GenreItem, user: dict = Depends(_require_auth)) -> dict[str, Any]:
    all_genres = _merged_genres(user["username"])
    if any(g["name"] == item.name for g in all_genres):
        raise HTTPException(status_code=409, detail=f"类型 '{item.name}' 已存在")
    try:
        return add_user_genre(user["username"], item.name, item.guidance, item.keywords)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/api/genres/{index}")
async def update_genre(index: int, item: GenreItem, user: dict = Depends(_require_auth)) -> dict[str, Any]:
    system = _load_genres()
    user_genres = get_user_genres(user["username"])
    system_len = len(system)

    if index < system_len:
        raise HTTPException(status_code=403, detail="系统默认类型不可修改")

    user_index = index - system_len
    all_user_names = {g["name"] for i, g in enumerate(user_genres) if i != user_index}
    if item.name in all_user_names:
        raise HTTPException(status_code=409, detail=f"类型 '{item.name}' 已存在")

    try:
        return update_user_genre(user["username"], user_index, item.name, item.guidance, item.keywords)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/genres/{index}")
async def delete_genre(index: int, user: dict = Depends(_require_auth)) -> dict[str, str]:
    system = _load_genres()
    system_len = len(system)

    if index < system_len:
        raise HTTPException(status_code=403, detail="系统默认类型不可删除")

    user_index = index - system_len
    try:
        deleted_name = delete_user_genre(user["username"], user_index)
        return {"deleted": deleted_name}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


# --- Conversion ---

@app.post("/api/convert", response_model=ConvertResponse)
@limiter.limit("5/minute")
async def convert(request: Request, req: ConvertRequest, user: dict = Depends(_require_auth)) -> ConvertResponse:
    text = req.text.strip()
    if len(text) < 100:
        raise HTTPException(status_code=400, detail="文本内容过短，至少需要 100 个字符")

    genre = req.genre.strip()
    genres = _load_genres()
    valid_names = {g["name"] for g in genres}
    if genre not in valid_names:
        genre = genres[0]["name"] if genres else "叙事"

    task_id = str(uuid.uuid4())
    db_create_conversion(task_id, user["username"], text, genre, req.title.strip())

    asyncio.create_task(_run_pipeline(task_id, text, genre, req.title.strip(), user["username"]))
    return ConvertResponse(task_id=task_id)


@app.get("/api/convert/{task_id}/progress")
async def progress(task_id: str, token: str = ""):
    _check_auth(token)

    conv = db_get_conversion(task_id)
    if not conv:
        raise HTTPException(status_code=404, detail="任务不存在")

    async def event_stream():
        while True:
            conv = db_get_conversion(task_id)
            if conv is None:
                break

            progress_data = conv.get("progress_json", {}) or {}
            status = conv.get("status", "pending")
            error = conv.get("error", "")

            event_data = {
                **progress_data,
                "status": status,
                "error": error,
            }

            yield f"data: {json.dumps(event_data, ensure_ascii=False)}\n\n"

            if status in ("completed", "failed"):
                break

            await asyncio.sleep(0.5)

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/api/convert/{task_id}/result")
async def result(task_id: str, user: dict = Depends(_require_auth)):
    conv = db_get_conversion(task_id)
    if not conv:
        raise HTTPException(status_code=404, detail="任务不存在")

    status = conv.get("status", "pending")
    if status in ("processing", "pending"):
        raise HTTPException(status_code=202, detail="任务尚未完成")

    if status == "failed":
        raise HTTPException(status_code=500, detail=f"任务失败: {conv.get('error', '未知错误')}")

    return {
        "yaml": conv.get("yaml_output", ""),
        "meta": conv.get("meta_json", {}),
    }


class RegenerateRequest(BaseModel):
    hints: str


@app.post("/api/convert/{task_id}/regenerate")
@limiter.limit("3/minute")
async def regenerate(request: Request, task_id: str, req: RegenerateRequest, user: dict = Depends(_require_auth)):
    conv = db_get_conversion(task_id)
    if not conv:
        raise HTTPException(status_code=404, detail="任务不存在")

    if conv.get("status") != "completed":
        raise HTTPException(status_code=400, detail="任务尚未完成，无法重新生成")

    intermediate = conv.get("intermediate_json", {}) or {}
    if not intermediate.get("original_text"):
        raise HTTPException(status_code=400, detail="无法重新生成，请重新转换")

    if not req.hints.strip():
        raise HTTPException(status_code=400, detail="请输入补充信息")

    original_text = intermediate["original_text"]
    genre = intermediate.get("genre", "叙事")
    title = intermediate.get("title", "")
    enriched_text = original_text + "\n\n【用户补充信息】\n" + req.hints.strip()

    db_update_conversion_status(task_id, "regenerating")
    db_update_conversion_progress(task_id, {
        "step": 0,
        "total": 7,
        "step_name": "初始化",
        "message": "正在根据补充信息重新生成...",
    })

    await _run_pipeline(task_id, enriched_text, genre, title, user["username"])

    conv = db_get_conversion(task_id)
    if conv:
        return {
            "yaml": conv.get("yaml_output", ""),
            "meta": conv.get("meta_json", {}),
        }
    return {}


# --- Conversion History ---

@app.get("/api/conversions")
async def list_conversions(user: dict = Depends(_require_auth)) -> list[dict[str, Any]]:
    return db_list_conversions(user["username"])


@app.delete("/api/conversions/{task_id}")
async def delete_conversion(task_id: str, user: dict = Depends(_require_auth)) -> dict[str, str]:
    deleted = db_delete_conversion(task_id, user["username"])
    if not deleted:
        raise HTTPException(status_code=404, detail="记录不存在")
    return {"deleted": task_id}


class EditYamlRequest(BaseModel):
    yaml: str


@app.put("/api/conversions/{task_id}")
async def edit_conversion(task_id: str, req: EditYamlRequest, user: dict = Depends(_require_auth)) -> dict[str, Any]:
    conv = db_get_conversion(task_id)
    if not conv:
        raise HTTPException(status_code=404, detail="任务不存在")
    if conv.get("status") != "completed":
        raise HTTPException(status_code=400, detail="只能编辑已完成的转换结果")

    updated = db_update_conversion_yaml(task_id, req.yaml, user["username"])
    if not updated:
        raise HTTPException(status_code=403, detail="无权编辑此记录")

    conv = db_get_conversion(task_id)
    return {
        "yaml": conv.get("yaml_output", ""),
        "meta": conv.get("meta_json", {}),
    }


# --- File Upload ---

@app.post("/api/upload")
@limiter.limit("10/minute")
async def upload_file(request: Request, file: UploadFile = File(...), user: dict = Depends(_require_auth)) -> dict[str, str]:
    ALLOWED_EXTENSIONS = {".txt", ".md", ".text", ".pdf", ".docx", ".doc"}
    ext = Path(file.filename or "").suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {ext}，仅支持 {', '.join(sorted(ALLOWED_EXTENSIONS))}")

    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过 5MB")

    text = ""

    if ext in {".txt", ".md", ".text"}:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content.decode("gbk")
            except UnicodeDecodeError:
                raise HTTPException(status_code=400, detail="无法识别文件编码，请使用 UTF-8 或 GBK 编码")

    elif ext == ".pdf":
        try:
            from io import BytesIO
            from PyPDF2 import PdfReader
            reader = PdfReader(BytesIO(content))
            parts: list[str] = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
            text = "\n".join(parts)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF 解析失败: {e}")

    elif ext == ".docx":
        try:
            from io import BytesIO
            from docx import Document
            doc = Document(BytesIO(content))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            text = "\n".join(parts)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX 解析失败: {e}")

    elif ext == ".doc":
        try:
            from io import BytesIO
            from docx import Document
            doc = Document(BytesIO(content))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            text = "\n".join(parts)
        except Exception:
            import re
            raw = content.decode("utf-8", errors="ignore")
            chunks = re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\w]{4,}', raw)
            if chunks:
                text = "\n".join(chunks)
            else:
                raise HTTPException(status_code=400, detail="DOC 文件解析失败，请转换为 DOCX 格式后重试")

    if not text.strip():
        raise HTTPException(status_code=400, detail="文件内容为空或无法提取文本")

    return {"filename": file.filename or "unknown", "text": text}


# --- Export ---

@app.get("/api/convert/{task_id}/export/pdf")
async def export_pdf_endpoint(task_id: str, token: str = ""):
    _check_auth(token)
    conv = db_get_conversion(task_id)
    if not conv or conv.get("status") != "completed":
        raise HTTPException(status_code=404, detail="转换任务不存在或未完成")

    yaml_str = conv.get("yaml_output", "")
    if not yaml_str:
        raise HTTPException(status_code=400, detail="无可用内容")

    try:
        pdf_bytes = export_pdf(yaml_str)
    except ExportError as e:
        raise HTTPException(status_code=400, detail=str(e))

    title = (conv.get("meta_json", {}) or {}).get("title", "script")
    filename = f"{title}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/convert/{task_id}/export/docx")
async def export_docx_endpoint(task_id: str, token: str = ""):
    _check_auth(token)
    conv = db_get_conversion(task_id)
    if not conv or conv.get("status") != "completed":
        raise HTTPException(status_code=404, detail="转换任务不存在或未完成")

    yaml_str = conv.get("yaml_output", "")
    if not yaml_str:
        raise HTTPException(status_code=400, detail="无可用内容")

    try:
        docx_bytes = export_docx(yaml_str)
    except ExportError as e:
        raise HTTPException(status_code=400, detail=str(e))

    title = (conv.get("meta_json", {}) or {}).get("title", "script")
    filename = f"{title}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# --- Health ---

@app.get("/api/health")
async def health() -> dict[str, str]:
    return {"status": "ok", "version": "0.4.0"}


# --- Serve frontend in production ---

FRONTEND_DIST = Path(__file__).parent.parent / "frontend" / "dist"
if FRONTEND_DIST.exists():
    app.mount("/", StaticFiles(directory=str(FRONTEND_DIST), html=True), name="frontend")
